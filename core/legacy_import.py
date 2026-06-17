from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from core.db import ATTACHMENT_DIR, IMPORT_DIR, add_workflow_event, df_query, json_dump, log_audit, make_ref, now_iso, run_insert, run_query

AMOUNT_RE = re.compile(r"(?:₦|NGN|N)?\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.\d{1,2})?|[0-9]+(?:\.\d{1,2})?)", re.I)
DATE_PATTERNS = [
    r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b",
    r"\b(\d{4}[/-]\d{1,2}[/-]\d{1,2})\b",
    r"\b(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{2,4})\b",
]

CATEGORY_KEYWORDS = {
    "Diesel/Fuel": ["diesel", "fuel", "petrol", "ago"],
    "Water": ["water"],
    "Office Supplies": ["a4", "paper", "rim", "ream", "office", "stationery", "desk", "chair"],
    "Generator Maintenance": ["generator", "gen", "100kva", "battery", "oil filter", "fuel filter"],
    "Vehicle Maintenance": ["vehicle", "hiace", "bus", "coaster", "tyre", "mechanic"],
    "Plumbing": ["plumbing", "pipe", "elbow", "tee", "tap"],
    "Welding/Fabrication": ["welding", "fabrication", "h-beam", "tank stand", "burglary proof", "iron"],
    "Grass Cutting": ["grass", "cutting", "mowing"],
    "Repairs/Maintenance": ["repair", "maintenance", "servicing", "service"],
    "Operational Purchases": ["operational", "operations", "purchase log"],
    "Staff Welfare": ["refreshment", "training", "welfare", "food"],
}

DOC_TYPE_KEYWORDS = [
    ("purchase log", "Purchase Log"),
    ("quotation", "Quotation"),
    ("quote", "Quotation"),
    ("requisition", "Requisition"),
    ("fuel", "Fuel Request"),
    ("maintenance", "Maintenance Request"),
    ("repair", "Repair Request"),
    ("report", "Report"),
    ("operational", "Operational Purchase"),
]


def safe_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._ -]+", "_", Path(name).name).strip() or "document.docx"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def parse_amount(value: Any) -> float:
    text = str(value or "").replace("\xa0", " ").strip()
    if not text:
        return 0.0
    # Prefer values close to currency markers if present, otherwise parse numbers.
    matches = AMOUNT_RE.findall(text)
    values = []
    for raw in matches:
        try:
            n = float(raw.replace(",", ""))
            if n >= 0:
                values.append(n)
        except Exception:
            pass
    return max(values) if values else 0.0


def parse_quantity(value: Any) -> float:
    text = str(value or "").strip()
    m = re.search(r"\d+(?:\.\d+)?", text.replace(",", ""))
    return float(m.group(0)) if m else 0.0


def normalize_header(text: str) -> str:
    t = re.sub(r"\s+", " ", str(text or "").strip().lower())
    t = t.replace("unite price", "unit price")
    if t in {"sn", "s/n", "sno", "no", "no."}:
        return "sn"
    if "date" in t:
        return "date"
    if "item" in t or "description" in t or "particular" in t:
        return "item"
    if "quant" in t or t == "qty":
        return "quantity"
    if "unit" in t and "price" in t:
        return "unit_price"
    if "total" in t and "price" in t or t == "amount" or "total" == t:
        return "total_price"
    if "status" in t:
        return "status_of_purchase"
    return t.replace(" ", "_")


def extract_docx_text_and_tables(docx_path: Path) -> tuple[str, list[list[list[str]]]]:
    document = Document(docx_path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
    table_text = []
    for table in tables:
        for row in table:
            table_text.append(" | ".join(row))
    return "\n".join(paragraphs + table_text), tables


def detect_department(original_path: str) -> str:
    upper = original_path.upper()
    if "CMOTD" in upper and "RACAM" in upper:
        return "CMOTD AND RACAM"
    if "CMOTD" in upper:
        return "CMOTD"
    if "RACAM" in upper:
        return "RACAM"
    parts = [p for p in Path(original_path).parts if p and p.upper() != "PROCUREMENT PROJECT"]
    return parts[0] if parts else "General"


def detect_doc_type(name: str, text: str) -> str:
    hay = f"{name}\n{text}".lower()
    for keyword, dtype in DOC_TYPE_KEYWORDS:
        if keyword in hay:
            return dtype
    return "Other"


def detect_category(name: str, text: str) -> str:
    hay = f"{name}\n{text}".lower()
    for category, words in CATEGORY_KEYWORDS.items():
        if any(w in hay for w in words):
            return category
    return "Other"


def extract_date(text: str, name: str = "") -> str | None:
    joined = f"{name}\n{text}"
    for pattern in DATE_PATTERNS:
        m = re.search(pattern, joined, flags=re.I)
        if not m:
            continue
        raw = m.group(1)
        for fmt in ["%m/%d/%Y", "%m/%d/%y", "%d/%m/%Y", "%d/%m/%y", "%Y/%m/%d", "%Y-%m-%d", "%d-%m-%Y", "%d-%m-%y", "%d %b %Y", "%d %B %Y", "%d %b %y", "%d %B %y"]:
            try:
                return datetime.strptime(raw, fmt).date().isoformat()
            except ValueError:
                continue
        return raw
    return None


def extract_title(text: str, file_name: str) -> str:
    lines = [re.sub(r"\s+", " ", l).strip() for l in text.splitlines() if l.strip()]
    for line in lines[:12]:
        if len(line) < 4:
            continue
        if re.fullmatch(r"[\d/\-\. ]+", line):
            continue
        if "total requisitioned amount" in line.lower():
            continue
        if len(line) <= 180:
            return line
    return Path(file_name).stem


def extract_total(text: str, line_items: list[dict[str, Any]]) -> float:
    candidates = []
    for pattern in [
        r"Total\s+Requisitioned\s+Amount\D{0,80}(?:₦|NGN|N)?\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.\d{1,2})?|[0-9]+(?:\.\d{1,2})?)",
        r"Total\s+Expenditure\D{0,80}(?:₦|NGN|N)?\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.\d{1,2})?|[0-9]+(?:\.\d{1,2})?)",
        r"Total\s+Allocated\D{0,80}(?:₦|NGN|N)?\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.\d{1,2})?|[0-9]+(?:\.\d{1,2})?)",
    ]:
        for m in re.finditer(pattern, text, flags=re.I):
            candidates.append(parse_amount(m.group(1)))
    item_total = sum(float(item.get("total_price") or 0) for item in line_items)
    if item_total > 0:
        candidates.append(item_total)
    return max(candidates) if candidates else 0.0


def parse_line_items(tables: list[list[list[str]]], category: str) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for table in tables:
        if not table or len(table) < 2:
            continue
        header_idx = None
        header_map = {}
        for i, row in enumerate(table[:3]):
            normalized = [normalize_header(c) for c in row]
            if any(h in normalized for h in ["item", "quantity", "unit_price", "total_price"]):
                header_idx = i
                header_map = {name: idx for idx, name in enumerate(normalized)}
                break
        if header_idx is None:
            continue
        for rno, row in enumerate(table[header_idx + 1 :], start=1):
            def get(key: str) -> str:
                idx = header_map.get(key)
                return row[idx] if idx is not None and idx < len(row) else ""
            item_name = get("item")
            if not item_name or item_name.strip().lower() in {"item", "total", "grand total"}:
                continue
            quantity = parse_quantity(get("quantity")) or 1.0
            unit_price = parse_amount(get("unit_price"))
            total_price = parse_amount(get("total_price"))
            if not total_price and unit_price and quantity:
                total_price = unit_price * quantity
            if not unit_price and total_price and quantity:
                unit_price = total_price / quantity if quantity else 0
            status = get("status_of_purchase")
            raw = {"row": row, "header_map": header_map}
            items.append({
                "row_number": rno,
                "item_name": re.sub(r"\s+", " ", item_name).strip()[:250],
                "description": item_name.strip(),
                "quantity": quantity,
                "unit_price": unit_price,
                "total_price": total_price,
                "category": category,
                "status_of_purchase": status,
                "raw_json": raw,
            })
    return items


def confidence_score(text: str, line_items: list[dict[str, Any]], total: float, likely_date: str | None) -> float:
    score = 0.25
    if text.strip():
        score += 0.15
    if line_items:
        score += 0.3
    if total > 0:
        score += 0.2
    if likely_date:
        score += 0.1
    return round(min(score, 0.98), 2)


def likely_vendor_from_text(text: str) -> str:
    # Real documents mostly use signatures and titles rather than explicit vendors. Keep this conservative.
    for pattern in [r"Vendor\s*[:\-]\s*(.+)", r"Supplier\s*[:\-]\s*(.+)", r"Quote\s+from\s+(.+)"]:
        m = re.search(pattern, text, flags=re.I)
        if m:
            return m.group(1).splitlines()[0][:120].strip()
    return ""


def insert_imported_document(parsed: dict[str, Any], user_id: int | None) -> tuple[int | None, str]:
    existing = df_query("SELECT id, linked_request_id FROM imported_legacy_documents WHERE original_path=? OR file_hash=?", (parsed["original_path"], parsed["file_hash"]))
    if not existing.empty:
        return int(existing.iloc[0]["id"]), "skipped_duplicate"

    doc_id = run_insert(
        """
        INSERT INTO imported_legacy_documents (
            source_zip_name, original_path, file_name, file_path, file_hash, document_type,
            department_project, title, likely_date, likely_vendor, total_amount, import_status,
            confidence, extracted_text, parsed_json, duplicate_warning, imported_by, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'Imported - Needs Review', ?, ?, ?, 0, ?, ?, ?)
        """,
        (
            parsed["source_zip_name"], parsed["original_path"], parsed["file_name"], parsed["file_path"], parsed["file_hash"],
            parsed["document_type"], parsed["department_project"], parsed["title"], parsed["likely_date"], parsed["likely_vendor"], parsed["total_amount"],
            parsed["confidence"], parsed["extracted_text"], json_dump(parsed), user_id, now_iso(), now_iso(),
        ),
    )

    for item in parsed["line_items"]:
        run_query(
            """
            INSERT INTO parsed_document_line_items (
                imported_doc_id, row_number, item_name, description, quantity, unit_price,
                total_price, category, status_of_purchase, raw_json, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                doc_id, item["row_number"], item["item_name"], item["description"], item["quantity"], item["unit_price"],
                item["total_price"], item["category"], item.get("status_of_purchase"), json_dump(item.get("raw_json", {})), now_iso(),
            ),
        )

    request_no = make_ref("LEG-PR")
    requester_id = user_id or 1
    request_id = run_insert(
        """
        INSERT INTO purchase_requests (
            request_no, requested_by, department_project, request_date, required_date, category,
            justification, priority, estimated_amount, vendor_preference, status, source_type,
            imported_doc_id, import_confidence, attachments_json, notes, approval_history_json,
            created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, 'Normal', ?, ?, 'Draft', 'Legacy Import', ?, ?, ?, ?, '[]', ?, ?)
        """,
        (
            request_no, requester_id, parsed["department_project"], parsed["likely_date"] or datetime.now().date().isoformat(),
            parsed["likely_date"], parsed["category"], parsed["title"], parsed["total_amount"], parsed["likely_vendor"],
            doc_id, parsed["confidence"], json_dump([parsed["file_path"]]), f"Imported from {parsed['original_path']}. Review before submission.", now_iso(), now_iso(),
        ),
    )
    for item in parsed["line_items"][:80]:
        run_query(
            """
            INSERT INTO purchase_request_items
            (request_id, item_name, description, quantity, unit_price, total, category, suggested_vendor, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                request_id, item["item_name"] or parsed["title"], item["description"], item["quantity"] or 1,
                item["unit_price"], item["total_price"], item["category"], parsed["likely_vendor"], now_iso(),
            ),
        )
    if not parsed["line_items"]:
        run_query(
            """
            INSERT INTO purchase_request_items
            (request_id, item_name, description, quantity, unit_price, total, category, suggested_vendor, created_at)
            VALUES (?, ?, ?, 1, ?, ?, ?, ?, ?)
            """,
            (request_id, parsed["title"], parsed["title"], parsed["total_amount"], parsed["total_amount"], parsed["category"], parsed["likely_vendor"], now_iso()),
        )
    run_query("UPDATE imported_legacy_documents SET linked_request_id=?, updated_at=? WHERE id=?", (request_id, now_iso(), doc_id))
    add_workflow_event("Imported Document", doc_id, "Imported", "Imported - Needs Review", parsed["original_path"], user_id)
    add_workflow_event("Purchase Request", request_id, "Created from legacy document", "Draft", parsed["title"], user_id)
    return doc_id, "imported"


def parse_docx_from_zip_member(zf: zipfile.ZipFile, member: str, source_zip_name: str, import_batch_dir: Path) -> dict[str, Any]:
    data = zf.read(member)
    fhash = sha256_bytes(data)
    relative_parts = [safe_name(p) for p in Path(member).parts if p and p != "PROCUREMENT PROJECT"]
    target_dir = import_batch_dir / Path(*relative_parts[:-1]) if len(relative_parts) > 1 else import_batch_dir
    target_dir.mkdir(parents=True, exist_ok=True)
    file_path = target_dir / relative_parts[-1]
    if not file_path.exists():
        file_path.write_bytes(data)

    text, tables = extract_docx_text_and_tables(file_path)
    category = detect_category(member, text)
    line_items = parse_line_items(tables, category)
    total = extract_total(text, line_items)
    likely_date = extract_date(text, Path(member).name)
    title = extract_title(text, Path(member).name)
    doc_type = detect_doc_type(member, text)
    confidence = confidence_score(text, line_items, total, likely_date)
    return {
        "source_zip_name": source_zip_name,
        "original_path": member,
        "file_name": Path(member).name,
        "file_path": str(file_path),
        "file_hash": fhash,
        "document_type": doc_type,
        "department_project": detect_department(member),
        "title": title,
        "likely_date": likely_date,
        "likely_vendor": likely_vendor_from_text(text),
        "category": category,
        "total_amount": total,
        "confidence": confidence,
        "extracted_text": text,
        "tables": tables,
        "line_items": line_items,
    }


def import_procurement_zip(zip_path: str | Path, user_id: int | None = None) -> dict[str, Any]:
    zip_path = Path(zip_path)
    batch_dir = IMPORT_DIR / f"import_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    batch_dir.mkdir(parents=True, exist_ok=True)
    summary = {"source_zip": zip_path.name, "imported": 0, "skipped": 0, "failed": 0, "partial": 0, "documents": []}
    with zipfile.ZipFile(zip_path) as zf:
        for member in zf.namelist():
            filename = Path(member).name
            if member.endswith("/"):
                continue
            if filename.startswith("~$") or filename.startswith("~WRL"):
                summary["skipped"] += 1
                run_query("INSERT INTO document_extraction_logs (source_zip_name, original_path, action, status, message, created_at) VALUES (?, ?, 'SKIP', 'Skipped', ?, ?)", (zip_path.name, member, "Temporary Word lock/temp file ignored", now_iso()))
                continue
            if not filename.lower().endswith(".docx"):
                summary["skipped"] += 1
                run_query("INSERT INTO document_extraction_logs (source_zip_name, original_path, action, status, message, created_at) VALUES (?, ?, 'SKIP', 'Skipped', ?, ?)", (zip_path.name, member, "Unsupported file type", now_iso()))
                continue
            try:
                parsed = parse_docx_from_zip_member(zf, member, zip_path.name, batch_dir)
                doc_id, status = insert_imported_document(parsed, user_id)
                if status == "imported":
                    summary["imported"] += 1
                else:
                    summary["skipped"] += 1
                if parsed["confidence"] < 0.55 or parsed["total_amount"] <= 0:
                    summary["partial"] += 1
                summary["documents"].append({"id": doc_id, "path": member, "status": status, "title": parsed["title"], "amount": parsed["total_amount"], "confidence": parsed["confidence"]})
                run_query("INSERT INTO document_extraction_logs (source_zip_name, original_path, action, status, message, created_at) VALUES (?, ?, 'IMPORT', ?, ?, ?)", (zip_path.name, member, status, f"{parsed['title']} | {parsed['total_amount']} | confidence {parsed['confidence']}", now_iso()))
            except Exception as exc:
                summary["failed"] += 1
                run_query("INSERT INTO document_extraction_logs (source_zip_name, original_path, action, status, message, created_at) VALUES (?, ?, 'IMPORT', 'Failed', ?, ?)", (zip_path.name, member, str(exc), now_iso()))
    log_audit("LEGACY_IMPORT", "ImportBatch", zip_path.name, summary, user_id)
    return summary


def import_uploaded_zip(uploaded_file, user_id: int | None = None) -> dict[str, Any]:
    if uploaded_file is None:
        return {"error": "No ZIP uploaded"}
    target = IMPORT_DIR / safe_name(uploaded_file.name)
    target.write_bytes(uploaded_file.getvalue())
    return import_procurement_zip(target, user_id)


def bundled_legacy_zip_path() -> Path:
    return Path(__file__).resolve().parents[1] / "sample_imports" / "PROCUREMENT PROJECT.zip"
